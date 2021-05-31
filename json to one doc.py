# -*- coding: utf-8 -*-
"""
Created on Mon May 31 15:14:44 2021

@author: Say My Name
"""

import pandas as pd
import numpy as np
import json
from docx import Document
def check_parent(file,arr,row,col):
    """this function will take children of json file, arr, row index and col index
    and return boolean value i.e it's checks whether the parent of a particular cell
    exists in that particular column"""
    for chk in range(row):
        #print(chk,file['header'][0],arr[chk][col][0])
        if file['header'][0]==arr[chk][col][0]:
            return True
        if file['header'][0]=="None" and arr[chk][col][0]==0:
            return True
    return False

def make_python_table(data):
    """this function makes a blueprint of table, i.e it's stores the table in
    the list"""
    rc=get_rows_col(data)
    #arr=[[[0]*2]*rc[1]]*rc[0]
    """creating list of dimension row*col*2
    at [row][col][0] contains id
    at [row][col][1] contains the text"""
    arr = [0] * rc[0]
    for i in range(rc[0]):
        arr[i] = [0] * rc[1]
        for z in range(rc[1]):
            arr[i][z]=[0]*2
    #filling first row of list
    d2=data['hasChildren'][0]
    j=0
    for child in d2['hasChildren']:
        cs=int(child['column_Span'])+j
        while j<cs:
            arr[0][j][0]=child['id']
            arr[0][j][1]=child['text']
            j+=1
    #filling all rows in the list
    d=data['hasChildren']
    row=1
    while row<rc[0]:
        col=0
        jcol=0
        while col<rc[1]:
            if check_parent(d[row]['hasChildren'][jcol],arr,row,col):
                col_span=int(d[row]['hasChildren'][jcol]['column_Span'])+col
                while col<col_span:
                    arr[row][col][0]=d[row]['hasChildren'][jcol]['id']
                    arr[row][col][1]=d[row]['hasChildren'][jcol]['text']
                    col+=1
                jcol+=1
                if jcol== len(d[row]['hasChildren']):
                    col=rc[1]
            else:
                col+=1
            
        row+=1
    """putting empty string where id of the cell is None"""
    for r1 in arr:
        for c1 in r1:
            if c1[0]==0:
                c1[1]=""
    return arr
        
def get_rows_col(data):
    """this function return dimension of the table"""
    r=len(data['hasChildren'])
    print(r)
    d1=data['hasChildren'][0]
    cnt=0
    for j in d1['hasChildren']:
        cnt=cnt+int(j['column_Span'])
    print(cnt)
    return [r,cnt]

def creating_table(data,document):
    rc=get_rows_col(data)
    arr=make_python_table(data)
    """creating table"""
    table = document.add_table(rows=1, cols=rc[1])
    table_row=table.rows[0].cells
    """copying text from arr to table"""
    for j_table in range(rc[1]):
        table_row[j_table].text=str(arr[0][j_table][1])
    i_table=1
    while i_table<rc[0]:
        table_row = table.add_row().cells
        for j_table in range(rc[1]):
            table_row[j_table].text=str(arr[i_table][j_table][1])
        i_table+=1
    """Now Merging the cell which has same id"""
    for i in range(rc[0]):
        j=1
        A=table.cell(i,j-1)
        while j<rc[1]:
            if arr[i][j-1][0]==arr[i][j][0]:
                a=table.cell(i,j)
                A=A.merge(a)
                A.text=str(arr[i][j][1])
            else:
                A=table.cell(i,j)
            j+=1

def creating_doc(a):
    """creating doc"""
    document=Document()
    document.add_heading('The TABLE',0)
    for filename in a:
        f= open(filename,)
        # returns JSON object as a dictionary
        data= json.load(f)
        creating_table(data, document)
        #closing file
        f.close()
        document.add_paragraph(' ')
    """output file name"""
    document.save('merge11.docx')


"""input json file"""
a=['New.json','tough.json','input3.json']
#f = open('New.json',)
# returns JSON object as a dictionary
#data = json.load(f)
  
# Iterating through the json
# list
#for i in data['people']:
    #print(i)
  
# Closing file
#f.close()  
#rc=get_rows_col(data)
creating_doc(a)
